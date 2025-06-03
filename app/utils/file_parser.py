import fitz  # PyMuPDF
import docx
# import pandas as pd # Commented out direct import
import os
import logging

# Try to import pandas for Excel support, but don't make it a hard requirement
try:
    import pandas as pd
    PANDAS_AVAILABLE = True
except ImportError:
    PANDAS_AVAILABLE = False
    pd = None # Ensure pd is defined even if not available
    logging.warning("Pandas library not found. Excel file (.xlsx, .xls) processing will be skipped.")

# Adjusted SUPPORTED_EXTENSIONS based on PANDAS_AVAILABLE
SUPPORTED_EXTENSIONS = [".pdf", ".docx", ".txt"]
if PANDAS_AVAILABLE:
    SUPPORTED_EXTENSIONS.extend([".xlsx", ".xls"])

def extract_text_from_pdf(file_path: str) -> str:
    """Extracts text content from a PDF file."""
    doc = fitz.open(file_path)
    text = ""
    for page_num in range(len(doc)):
        page = doc.load_page(page_num)
        text += page.get_text()
    doc.close()
    return text

def extract_text_from_docx(file_path: str) -> str:
    """Extracts text content from a .docx Word file."""
    doc = docx.Document(file_path)
    text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
    return text

def extract_text_from_excel(file_path: str) -> str:
    """Extracts text content from an Excel file (.xlsx, .xls).
       Converts each sheet to a CSV-like string representation.
       Requires pandas to be installed."""
    if not PANDAS_AVAILABLE:
        logging.warning(f"Skipping Excel file {file_path} as pandas library is not available.")
        return ""
    
    xls = pd.ExcelFile(file_path)
    text = ""
    for sheet_name in xls.sheet_names:
        df = xls.parse(sheet_name)
        text += f"Sheet: {sheet_name}\n{df.to_string(index=False)}\n\n"
    return text.strip()

def extract_text_from_txt(file_path: str) -> str:
    """Extracts text content from a .txt file."""
    with open(file_path, "r", encoding="utf-8") as f:
        return f.read()

def load_and_extract_text(file_path: str) -> str:
    """Loads a supported file and extracts its text content."""
    if not file_path:
        logging.info("No file path provided for text extraction.")
        return ""
    
    try:
        # Ensure the path is absolute or correctly relative to the project root.
        # This logic might need adjustment based on your execution context (e.g., where run.py is)
        if not os.path.isabs(file_path):
            # This assumes file_parser.py is in app/utils/ and services are in app/services/
            # and the paths in .env are relative to the project root.
            project_root = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", ".."))
            file_path = os.path.join(project_root, file_path)

        if not os.path.exists(file_path):
            logging.warning(f"Knowledge base file not found at: {file_path}")
            return ""

        _, extension = os.path.splitext(file_path.lower())

        logging.info(f"Attempting to load and extract text from: {file_path} (extension: {extension})")

        if extension == ".pdf":
            content = extract_text_from_pdf(file_path)
        elif extension == ".docx":
            content = extract_text_from_docx(file_path)
        elif extension == ".xlsx" or extension == ".xls":
            content = extract_text_from_excel(file_path)
        elif extension == ".txt":
            content = extract_text_from_txt(file_path)
        else:
            logging.warning(f"Unsupported file type: {extension} for file {file_path}. Only {SUPPORTED_EXTENSIONS} are supported.")
            return ""
        
        logging.info(f"Successfully extracted text from {file_path} ({len(content)} chars).")
        return content

    except Exception as e:
        logging.error(f"Error loading or parsing file {file_path}: {e}")
        return ""

def load_knowledge_from_directory(directory_path: str) -> str:
    """Loads and concatenates text from all supported files in a directory."""
    if not directory_path:
        logging.info("No directory path provided for knowledge base loading.")
        return ""

    # Resolve directory path relative to project root if necessary
    # This logic assumes file_parser.py is in app/utils/
    # and the paths in .env are relative to the project root.
    if not os.path.isabs(directory_path):
        project_root = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", ".."))
        directory_path = os.path.join(project_root, directory_path)

    if not os.path.isdir(directory_path):
        logging.warning(f"Knowledge base directory not found or is not a directory: {directory_path}")
        return ""

    all_text_content = []
    logging.info(f"Scanning directory for knowledge base files: {directory_path}")

    for filename in sorted(os.listdir(directory_path)): # Sort for consistent order
        file_path = os.path.join(directory_path, filename)
        if os.path.isfile(file_path):
            _, extension = os.path.splitext(filename.lower())
            if extension in SUPPORTED_EXTENSIONS:
                logging.info(f"Found supported file in directory: {filename}")
                # Use the existing load_and_extract_text to get content from this individual file
                # Note: load_and_extract_text already handles relative path resolution if called with a relative path
                # but here file_path is already absolute or correctly resolved relative to directory_path
                text = load_and_extract_text(file_path) 
                if text:
                    # Add a separator or filename marker for clarity if desired
                    all_text_content.append(f"--- Content from: {filename} ---\n{text}\n\n")
            else:
                logging.debug(f"Skipping unsupported file in directory: {filename}")
    
    if not all_text_content:
        logging.info(f"No supported files found or no content extracted from directory: {directory_path}")
        return ""
    
    concatenated_content = "".join(all_text_content)
    logging.info(f"Consolidated knowledge from directory {directory_path} ({len(concatenated_content)} chars from {len(all_text_content)} files).")
    return concatenated_content

if __name__ == '__main__':
    # Create dummy files for testing in a temporary 'test_data' directory
    # This assumes you run: python app/utils/file_parser.py from the project root.
    test_data_dir = "test_data_parser"
    os.makedirs(test_data_dir, exist_ok=True)

    txt_file = os.path.join(test_data_dir, "sample.txt")
    docx_file = os.path.join(test_data_dir, "sample.docx")
    # pdf_file = os.path.join(test_data_dir, "sample.pdf") # Requires a valid PDF to be created
    excel_file = os.path.join(test_data_dir, "sample.xlsx")

    with open(txt_file, "w") as f:
        f.write("This is a test text file.")

    doc = docx.Document()
    doc.add_paragraph("This is a test DOCX file.")
    doc.save(docx_file)
    
    if PANDAS_AVAILABLE:
        df_test = pd.DataFrame({'Col1': [1, 2], 'Col2': ['A', 'B']})
        with pd.ExcelWriter(excel_file, engine='openpyxl') as writer:
            df_test.to_excel(writer, sheet_name='Sheet1', index=False)
        print(f"\n--- Testing XLSX: {excel_file} --- (if pandas is installed)")
        print(load_and_extract_text(excel_file))
    else:
        print("\n--- Skipping XLSX test as pandas is not installed ---")
    # print(f"\n--- Testing PDF (requires sample.pdf): {pdf_file} ---")
    # print(load_and_extract_text(pdf_file))

    # --- Test for load_knowledge_from_directory ---
    print("\n\n--- Testing Directory Loading ---")
    test_kb_dir = "test_knowledge_dir"
    os.makedirs(test_kb_dir, exist_ok=True)
    # Create some dummy files in the test_kb_dir (relative to where this script is run)
    # For testing, ensure this script is run from the project root, or adjust paths.
    # If running `python app/utils/file_parser.py`, paths will be relative to `python-whatsapp-bot`

    txt_file_in_dir = os.path.join(test_kb_dir, "info1.txt")
    docx_file_in_dir = os.path.join(test_kb_dir, "report.docx")
    unsupported_file = os.path.join(test_kb_dir, "archive.zip")

    with open(txt_file_in_dir, "w") as f:
        f.write("This is info from info1.txt.")
    
    doc_dir = docx.Document()
    doc_dir.add_paragraph("This is a DOCX report for the knowledge base.")
    doc_dir.save(docx_file_in_dir)

    with open(unsupported_file, "w") as f:
        f.write("This is a zip file and should be ignored.")

    # Test loading from the directory (path relative to project root if script run from there)
    # If running `app/utils/file_parser.py` directly, the path should be relative to its location,
    # or better, use absolute paths for testing or ensure .env paths are correct for app execution.
    # For the purpose of this __main__ block, we assume the test_kb_dir is created in the CWD.
    knowledge_from_dir = load_knowledge_from_directory(test_kb_dir)
    print(f"Knowledge loaded from directory '{test_kb_dir}':\n{knowledge_from_dir}")

    # Clean up dummy directory and files
    # import shutil
    # shutil.rmtree(test_kb_dir) 
    # shutil.rmtree("test_data_parser") # Also clean up single file test dir if needed
    print(f"\nNote: Test files and directory created in ./{test_kb_dir}/ and ./test_data_parser/ . You may want to manually delete them or uncomment cleanup.") 